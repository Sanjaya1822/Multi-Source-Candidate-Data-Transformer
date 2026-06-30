"""
Resume Adapter -- Hybrid Extraction Pipeline.

Pipeline: text extraction -> section detection -> rule-based parsing
          -> NLP/NER (spaCy if available, heuristic fallback) -> normalization

Supports: PDF (pdfplumber), DOCX (python-docx), TXT, JSON fixture.

Sections detected: Contact, Summary, Skills, Experience, Projects,
                   Education, Certifications, Awards, Languages.

Trust score: 0.78 (text), 0.90 (JSON fixture)
"""
from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Any, Dict, List, Optional

from data_transformer.adapters.base import SourceAdapter
from data_transformer.schema.canonical import RawRecord, Source

# ---------------------------------------------------------------------------
# Compiled regex patterns
# ---------------------------------------------------------------------------

EMAIL_RE = re.compile(
    r"\b[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}\b"
)
PHONE_RE = re.compile(
    r"(?<!\d)"
    r"(\+?(?:\d{1,3}[\s\-\.]?)?\s*\(?\d{3,5}\)?[\s\-\.]?\d{3,5}[\s\-\.]?\d{3,5})"
    r"(?!\d)"
)
LINKEDIN_RE = re.compile(
    r"https?://(?:www\.)?linkedin\.com/in/[\w\-]+/?", re.I
)
GITHUB_PROFILE_RE = re.compile(
    r"https?://(?:www\.)?github\.com/([\w\-]+)(?:/[\w\-]+)?/?", re.I
)
URL_RE = re.compile(r"https?://[^\s,)\]>\"]+", re.I)

DATE_RE = re.compile(
    r"\b(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?"
    r"|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)"
    r"\.?\s+\d{4}\b"
    r"|\b\d{1,2}[/\-]\d{4}\b"
    r"|\b\d{4}[/\-]\d{2}\b"
    r"|\b\d{4}\b",
    re.I,
)

DATE_RANGE_RE = re.compile(
    r"((?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?"
    r"|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)"
    r"\.?\s+\d{4}|\d{1,2}[/\-]\d{4}|\d{4}[/\-]\d{2}|\d{4})"
    r"\s*(?:\u2013|\u2014|-{1,2}|to)\s*"
    r"((?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?"
    r"|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)"
    r"\.?\s+\d{4}|\d{1,2}[/\-]\d{4}|\d{4}[/\-]\d{2}|\d{4}|[Pp]resent|[Cc]urrent|[Nn]ow)",
    re.I,
)

SECTION_HEADER_RE = re.compile(
    r"^[ \t]*(?P<title>"
    # Experience variants
    r"(?:WORK\s+)?EXPERIENCE|PROFESSIONAL\s+EXPERIENCE|EMPLOYMENT(?:\s+HISTORY)?"
    r"|WORK\s+HISTORY|WORK\s+EXPERIENCE"
    r"|INTERNSHIP(?:\s+EXPERIENCE)?|INTERNSHIPS?(?:\s+&\s+EXPERIENCE)?"
    r"|INDUSTRIAL\s+TRAINING|TRAINING(?:\s+EXPERIENCE)?"
    r"|WORK\s+EXPERIENCE\s+&\s+INTERNSHIPS?"
    # Education variants
    r"|EDUCATION(?:AL)?(?:\s+BACKGROUND|\s+DETAILS|\s+QUALIFICATIONS?)?"
    r"|ACADEMIC(?:\s+BACKGROUND|\s+QUALIFICATIONS?|\s+DETAILS)?"
    r"|QUALIFICATION(?:S)?"
    # Skills variants
    r"|(?:TECHNICAL\s+|KEY\s+|PROFESSIONAL\s+|CORE\s+)?SKILLS?"
    r"|COMPETENCIES|CORE\s+COMPETENCIES|EXPERTISE|TECHNOLOGIES"
    r"|TOOLS(?:\s+&\s+TECHNOLOGIES)?|TECH(?:NICAL)?\s+STACK"
    # Projects
    r"|(?:PERSONAL\s+|SIDE\s+|ACADEMIC\s+|KEY\s+)?PROJECTS?"
    # Certs
    r"|CERTIFICATIONS?|LICENSES?(?:\s*(?:&|AND)?\s*CERTIFICATIONS?)?"
    r"|COURSES?(?:\s+&\s+CERTIFICATIONS?)?"
    # Summary / objective
    r"|CAREER\s+OBJECTIVE|CAREER\s+SUMMARY|PROFESSIONAL\s+SUMMARY"
    r"|SUMMARY(?:\s+OF\s+QUALIFICATIONS?)?|OBJECTIVE|PROFILE|ABOUT\s+ME"
    # Contact
    r"|CONTACT(?:\s+INFORMATION|\s+DETAILS)?|PERSONAL\s+DETAILS?"
    r"|PERSONAL\s+INFORMATION"
    # Awards / misc
    r"|AWARDS?|HONORS?|ACHIEVEMENTS?|ACCOMPLISHMENTS?"
    r"|LANGUAGES?|HOBBIES?|INTERESTS?|EXTRA\s+CURRICULAR"
    r"|EXTRA[\s\-]CURRICULAR(?:\s+ACTIVITIES?)?"
    r"|PUBLICATIONS?|VOLUNTEERING?|REFERENCES?|DECLARATION"
    r")[ \t]*:?[ \t]*$",
    re.I | re.M,
)

SKILL_LABEL_RE = re.compile(
    r"^(?:&\s*)?(?:"
    r"Programming\s+Languages?|Languages?|Frameworks?\s*(?:&\s*Libraries?)?|Libraries?"
    r"|Tools?|Databases?|Platforms?|Technologies|Cloud(?:\s*&\s*DevOps)?|DevOps"
    r"|Testing|Others?|&\s*Libraries?|&\s*Frameworks?|&\s*DevOps|&\s*Tools?"
    r"|Soft\s+Skills?|Operating\s+Systems?|Methodologies?|Concepts?"
    r")"
    r"[\s:,\-]+",
    re.I,
)

LOCATION_RE = re.compile(
    r"\b([A-Z][a-z]+(?:[ \t\-][A-Z][a-z]+)*"
    r",\s*(?:[A-Z]{2}|[A-Z][a-z]+(?:[ \t]+[A-Z][a-z]+)*))"
    r"(?:\s*,\s*(?:USA?|U\.S\.A?|United\s+States|India|UK|Canada|Australia))?\b"
)

DEGREE_RE = re.compile(
    r"\b(B\.?S\.?|B\.?A\.?|B\.?E\.?|B\.?Tech\.?|M\.?S\.?|M\.?A\.?|M\.?E\.?"
    r"|M\.?Tech\.?|MBA|Ph\.?D\.?|M\.?Phil\.?|Associate|Bachelor|Master|Doctor(?:ate)?"
    r"|HSC|SSLC|B\.?Sc|M\.?Sc|B\.?Com)"
    r"(?:\s+(?:of|in|of\s+Science|of\s+Arts|of\s+Engineering|of\s+Technology))?\b",
    re.I,
)

CERT_RE = re.compile(
    r"(?:AWS|GCP|Azure|Google|Cisco|Oracle|Microsoft|CompTIA|PMI|Scrum|"
    r"Certified|Certificate|Certification|License|CCNA|CCNP|CPA|CFA|PMP|"
    r"Associate|Professional|Expert|Practitioner)[^\n]{0,120}",
    re.I,
)

ACTION_VERBS_RE = re.compile(
    r"\b(led|built|reduced|integrated|deployed|managed|designed|developed|"
    r"implemented|created|improved|optimized|launched|shipped|maintained|"
    r"collaborated|coordinated|delivered|architected|migrated|automated|"
    r"applied|learnt|utilized|used|worked)\b",
    re.I,
)


# ---------------------------------------------------------------------------
# Section detection
# ---------------------------------------------------------------------------

def _canonical_section(raw: str) -> str:
    r = raw.strip().upper()
    # Experience / Internship — all go into "experience"
    if re.search(r"EXPERIENCE|EMPLOYMENT|WORK\s+HISTORY", r):
        return "experience"
    if re.search(r"INTERNSHIP|TRAINING|INDUSTRIAL\s+TRAINING", r):
        return "experience"
    # Education
    if re.search(r"EDUCATION|ACADEMIC|QUALIFICATION", r):
        return "education"
    # Skills — only if the heading itself is a skills heading (not "Work Experience")
    if re.search(r"\bSKILL|TECHNOLOG|COMPETENC|EXPERTISE|TECH\s+STACK|TOOLS", r):
        return "skills"
    # Projects
    if re.search(r"PROJECT", r):
        return "projects"
    # Certifications
    if re.search(r"CERTIF|LICENSE|COURSE", r):
        return "certifications"
    # Summary / objective
    if re.search(r"CAREER\s+OBJECTIVE|CAREER\s+SUMMARY|SUMMARY|OBJECTIVE|PROFILE|ABOUT|DECLARATION", r):
        return "summary"
    # Contact
    if re.search(r"CONTACT|PERSONAL\s+(?:DETAIL|INFO)", r):
        return "contact"
    # Awards
    if re.search(r"AWARD|HONOR|ACHIEVEMENT|ACCOMPLISHMENT", r):
        return "awards"
    if re.search(r"LANGUAGE", r):
        return "languages"
    if re.search(r"HOBBY|INTEREST|EXTRA|VOLUNTEER|PUBLICATION|REFERENCE", r):
        return "other"
    return "other"


def _split_sections(text: str) -> Dict[str, str]:
    sections: Dict[str, str] = {}
    lines = text.splitlines()
    current_name = "header"
    current_lines: List[str] = []

    for line in lines:
        m = SECTION_HEADER_RE.match(line)
        if m:
            body = "\n".join(current_lines).strip()
            if body:
                sections[current_name] = (
                    sections.get(current_name, "") + "\n" + body
                ).strip()
            current_name = _canonical_section(m.group("title"))
            current_lines = []
        else:
            current_lines.append(line)

    body = "\n".join(current_lines).strip()
    if body:
        sections[current_name] = (
            sections.get(current_name, "") + "\n" + body
        ).strip()

    return sections


# ---------------------------------------------------------------------------
# Date helpers
# ---------------------------------------------------------------------------

def _parse_date(raw: str) -> Optional[str]:
    if not raw:
        return None
    raw = raw.strip()
    if re.match(r"^(present|current|now)$", raw, re.I):
        return None
    try:
        from dateutil import parser as dp
        dt = dp.parse(raw, default=dp.parse("2000-01-01"))
        return dt.strftime("%Y-%m")
    except Exception:
        m = re.search(r"\b(\d{4})\b", raw)
        return f"{m.group(1)}-01" if m else None


# ---------------------------------------------------------------------------
# Contact extraction
# ---------------------------------------------------------------------------

def _spacy_available() -> bool:
    try:
        import spacy
        spacy.load("en_core_web_sm")
        return True
    except Exception:
        return False


def _extract_name(header_text: str, emails: List[str], phones: List[str]) -> str:
    # Try spaCy NER first
    try:
        import spacy
        nlp = spacy.load("en_core_web_sm")
        doc = nlp(header_text[:500])
        for ent in doc.ents:
            if ent.label_ == "PERSON" and 2 <= len(ent.text.split()) <= 5:
                return ent.text.strip()
    except Exception:
        pass

    # Heuristic: first short non-contact line
    noise = set(emails + phones)
    for line in header_text.splitlines()[:10]:
        line = line.strip()
        if not line:
            continue
        # Reject lines that are clearly not a name
        if (any(n in line for n in noise)
                or re.search(r"https?://|@|\|", line, re.I)
                or re.search(r"\blinkedin\b|\bgithub\b|\bresume\b|\bcv\b", line, re.I)
                or len(line) > 60
                or re.match(r"^\d", line)
                or SECTION_HEADER_RE.match(line)):
            continue
        # Reject section heading phrases (e.g. "Career Objective", "Technical Skills")
        if re.search(
            r"\b(objective|summary|profile|skills?|experience|education|"
            r"contact|certif|project|overview|introduction|career|professional)\b",
            line, re.I
        ):
            continue
        # Must look like a name: 2–5 words, mostly letters
        words = line.split()
        if 2 <= len(words) <= 5 and not re.search(r"\d", line):
            if all(re.match(r"^[A-Za-z\-\'\.]+$", w) for w in words):
                # Ensure at least one word has a capital letter or all caps
                if any(re.match(r"[A-Z]", w) for w in words):
                    return line.title() if line.isupper() else line
    return ""


# Patterns that indicate a string is NOT a location (education/institution junk)
_LOCATION_REJECT_RE = re.compile(
    r"\b(?:university|college|institute|campus|school|academy|polytechnic|"
    r"b\.?tech|b\.?e|m\.?tech|m\.?s|b\.?sc|cgpa|gpa|engineering|technology|"
    r"science|arts|commerce|management)\b",
    re.I,
)


def _extract_location(text: str) -> str:
    # 1. Explicit label
    m = re.search(
        r"(?:location|address|city|based\s+in|located\s+in)\s*[:\-]?\s*"
        r"([A-Za-z][A-Za-z\s\-,\.]+?)(?:\n|$)",
        text, re.I
    )
    if m:
        candidate = m.group(1).strip().rstrip(",.")
        if not _LOCATION_REJECT_RE.search(candidate):
            return candidate

    # 2. Pattern match "City, State" — only in header (first 400 chars)
    for m in LOCATION_RE.finditer(text[:400]):
        candidate = m.group(0).strip()
        if not _LOCATION_REJECT_RE.search(candidate):
            return candidate

    # 3. Fallback: Check for known cities/states if no comma
    fallback_re = re.compile(r"\b(New York|San Francisco|Los Angeles|Chicago|Seattle|Boston|Austin|London|Bangalore|Chennai|Hyderabad|Pune|Mumbai|Delhi|Toronto|Vancouver|Sydney|Melbourne|Coimbatore)\b", re.I)
    for m in fallback_re.finditer(text[:400]):
        return m.group(1).title()

    return ""


def _extract_contact(header_text: str, full_text: str) -> Dict[str, Any]:
    scan = header_text + "\n" + full_text[:2000]

    emails = sorted({
        e for e in EMAIL_RE.findall(full_text)
        if not re.search(r"\.(png|jpg|svg|gif)$", e, re.I)
    })

    # Phones: deduplicate by digit sequence
    seen_digits: set = set()
    phones: List[str] = []
    for p in PHONE_RE.findall(full_text[:3000]):
        digits = re.sub(r"\D", "", p)
        if len(digits) >= 10 and digits not in seen_digits:
            seen_digits.add(digits)
            phones.append(p.strip())

    linkedin_urls = sorted({u.rstrip("/") for u in LINKEDIN_RE.findall(full_text)})
    # GitHub: only profile-level (username only, no deep repo paths)
    github_urls: List[str] = []
    seen_gh: set = set()
    for m in GITHUB_PROFILE_RE.finditer(full_text):
        username = m.group(1).lower()
        url = f"https://github.com/{m.group(1)}"
        if username not in seen_gh:
            seen_gh.add(username)
            github_urls.append(url)

    links: List[Dict] = (
        [{"url": u, "type": "linkedin"} for u in linkedin_urls]
        + [{"url": u, "type": "github"}  for u in github_urls]
    )
    # Portfolio / personal site from header only
    for url in URL_RE.findall(header_text):
        url = url.rstrip(".,;)")
        if not any(d in url for d in ("linkedin.com", "github.com")):
            links.append({"url": url, "type": "portfolio"})

    full_name = _extract_name(header_text, emails, phones)
    location  = _extract_location(scan)

    return {
        "full_name": full_name,
        "emails":    emails,
        "phones":    phones,
        "links":     links,
        "location":  location,
    }


# ---------------------------------------------------------------------------
# Skills extraction
# ---------------------------------------------------------------------------

def _extract_skills(skills_text: str, _unused: str) -> List[str]:
    """
    Extract skills with label stripping and deduplication.

    Handles:
      - "Python, React, Node.js"
      - "Programming Languages: Java, Python"
      - "& Libraries: FastAPI, Django"
      - Bullet / indented lists

    Processes line-by-line to strip category labels correctly.
    Hard filters prevent experience/project text from bleeding in.
    """
    if not skills_text.strip():
        return []

    seen: Dict[str, str] = {}
    skills: List[str] = []

    for line in skills_text.splitlines():
        line = line.strip()
        # Strip bullet chars and leading & from line start
        line = re.sub(r"^[\u2022\u00b7\u25aa\u25b8\u25e6\-\u2013\s&]+", "", line)
        if not line:
            continue
        # Strip category label prefix
        line = SKILL_LABEL_RE.sub("", line).strip()
        if not line:
            continue
        # Split by common delimiters
        for item in re.split(r"[,;|]", line):
            # Strip surrounding punctuation and whitespace
            item = item.strip()
            item = re.sub(r'^[\u2013\u2014\u2022\u00b7\u25aa\u25b8\u25e6"\']+', "", item)
            item = re.sub(r'[\u2013\u2014\u2022\u00b7\u25aa\u25b8\u25e6"\']+$', "", item)
            item = item.strip()
            if not item:
                continue
            # Hard length filter
            if len(item) > 40:
                continue
            # Must contain letters
            if not re.search(r"[A-Za-z]", item):
                continue
            # Skip date-like items
            if re.search(
                r"\b(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)\b",
                item, re.I
            ):
                continue
            # Skip action-verb sentences (experience text bleeding in)
            if ACTION_VERBS_RE.search(item):
                continue
            # Skip URLs
            if re.search(r"https?://", item, re.I):
                continue
            # Skip sentence fragments (contain period = end of sentence or too many words)
            if (item.endswith(".") and " " in item) or len(item.split()) > 5:
                continue
            # Skip document titles and filenames
            if re.search(r"\.(pdf|docx?|txt)$|\b(resume|cv|document|profile)\b", item, re.I):
                continue
            # Skip date ranges (e.g. "June 2023 – August 2023")
            if DATE_RANGE_RE.search(item):
                continue
            # Skip lines that look like section headings
            if SECTION_HEADER_RE.match(item) or re.match(r"^(about|summary|objective|experience|education)$", item, re.I):
                continue
            # Skip items that look like job titles (contain Intern/Engineer/Manager etc.)
            if re.search(
                r"\b(intern|internship|engineer|developer|analyst|manager|coordinator|"
                r"architect|scientist|consultant|officer|executive|director|"
                r"associate|assistant|trainee|specialist|lead|head\s+of)\b",
                item, re.I
            ):
                continue
            # Skip items that look like company/org names.
            # A company name has a corporate suffix like Corp/Ltd/LLC/Hub
            # AND doesn't look like a tech term.
            # We check: has corp suffix, is not a known tech phrase, has no '/'
            _corp_suffix = re.search(
                r"\b(corporation|pvt\.?\s*ltd\.?|private\s+limited|"
                r"llc\b|inc\.\B|"
                r"iON\s+digital|digital\s+learning\s+hub)\b",
                item, re.I
            )
            if _corp_suffix:
                continue
            # Skip items that look like city names used as skills
            # (simple heuristic: single capitalized word that matches location RE)
            if LOCATION_RE.match(item):
                continue
            key = item.lower()
            if key not in seen:
                seen[key] = item
                skills.append(item)

    return skills


# ---------------------------------------------------------------------------
# Project tech keyword extraction (separate from full skills extraction)
# ---------------------------------------------------------------------------

def _extract_project_tech(proj_text: str) -> List[str]:
    """
    Extract only explicit technology keywords from project blocks.
    Only reads 'Technologies:', 'Stack:', 'Built with:' label lines.
    Requires a colon after the keyword to avoid matching description sentences.
    """
    if not proj_text.strip():
        return []

    tech: List[str] = []
    seen: set = set()

    for line in proj_text.splitlines():
        line = line.strip()
        # Must have a colon — "Technologies: Python, Kafka" not "built using Kafka"
        tm = re.match(
            r"(?:technologies|tech(?:nology)?|built\s+with|stack|tools)"
            r"\s*:\s*(.+)",
            line, re.I
        )
        if tm:
            for item in re.split(r"[,;|]", tm.group(1)):
                item = item.strip()
                if item and len(item) <= 40 and re.search(r"[A-Za-z]", item):
                    key = item.lower()
                    if key not in seen:
                        seen.add(key)
                        tech.append(item)
    return tech


# ---------------------------------------------------------------------------
# Experience extraction
# ---------------------------------------------------------------------------

def _is_job_title(text: str) -> bool:
    """Return True if text looks like a job title."""
    return bool(re.search(
        r"\b(engineer|developer|intern|manager|analyst|designer|lead|director|"
        r"architect|scientist|consultant|specialist|officer|coordinator|associate|"
        r"executive|trainee|programmer|tester|devops|sre|sde|backend|frontend|"
        r"fullstack|full.stack|software|web|data|ml|ai|cloud|security)\b",
        text, re.I
    ))


def _strip_city_suffix(company: str) -> str:
    """Remove city/location suffix from company: "Zoho Corp, Chennai" -> "Zoho Corp"."""
    # Strip trailing ", CityName" or ", CityName, Country" patterns
    cleaned = re.sub(r",\s*[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*(?:,\s*[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)?\s*$", "", company).strip()
    return cleaned if cleaned else company


def _extract_experience(exp_text: str) -> List[Dict[str, Any]]:
    """
    Parse structured experience entries.

    Handles all resume formats:
      A. "Title | Company    dates"          (pipe-separated, senior format)
      B. "Title\nCompany, City\ndates"      (multi-line, fresher format)
      C. "Company\nTitle\ndates"            (company-first format)
      D. Bullet-only blocks                  (descriptions only)

    Internship sections and work experience sections both map here.
    City names are stripped from company fields.
    """
    if not exp_text.strip():
        return []

    entries: List[Dict[str, Any]] = []
    
    # Pre-process dense lists
    et = exp_text.strip()
    if "\n\n" not in et:
        lines = et.splitlines()
        new_lines = []
        for i, l in enumerate(lines):
            # Split before lines that have both a date and a title/company separator (single-line headers)
            if i > 0 and DATE_RANGE_RE.search(l) and re.search(r"\s+(?:\||\bat\b)\s+|,\s*", l) and len(l) < 100:
                new_lines.append("\n" + l)
            else:
                new_lines.append(l)
        et = "\n".join(new_lines)
        
    # Split on blank lines — each block is one job
    blocks = re.split(r"\n\s*\n", et)

    for block in blocks:
        block = block.strip()
        if not block:
            continue

        all_lines = [l.strip() for l in block.splitlines() if l.strip()]
        if not all_lines:
            continue

        entry: Dict[str, Any] = {
            "company": None, "title": None,
            "start_date": None, "end_date": None,
            "is_current": False, "description": None,
        }

        # ── 1. Find and remove date range ─────────────────────────────────
        dm = DATE_RANGE_RE.search(block)
        if dm:
            entry["start_date"] = _parse_date(dm.group(1))
            end_raw = dm.group(2)
            if re.match(r"^(present|current|now)$", end_raw, re.I):
                entry["is_current"] = True
            else:
                entry["end_date"] = _parse_date(end_raw)

        # ── 2. Separate header lines from bullet/description lines ─────────
        header_lines = []
        desc_lines   = []
        for line in all_lines:
            is_bullet = bool(re.match(r"^[\u2022\u25aa\u25b8\-\*]", line))
            is_date   = bool(DATE_RANGE_RE.search(line))
            is_long   = len(line) > 75
            is_sent   = bool(re.search(r"[.!?]\s", line)) or (len(line) > 40 and ACTION_VERBS_RE.search(line))
            is_cont   = bool(re.match(r"^[a-z]", line.lstrip()))
            
            # If it's just a date line with nothing else, it's part of the header
            if is_bullet or is_long or is_sent or is_cont:
                desc_lines.append(line)
            else:
                header_lines.append(line)

        # ── 3. Parse title + company from header lines ─────────────────────
        if header_lines:
            first = header_lines[0]

            # Format A: "Title | Company" or "Title at Company" or "Title, Company"
            pipe_sep = re.search(r"\s+(?:\||\bat\b)\s+|,\s*", first, re.I)
            if pipe_sep:
                left  = first[:pipe_sep.start()].strip()
                right = first[pipe_sep.end():].strip()
                # Strip dates from right if any
                dm = DATE_RANGE_RE.search(right)
                if dm:
                    right = right[:dm.start()].strip()
                
                if _is_job_title(left):
                    entry["title"]   = left
                    entry["company"] = _strip_city_suffix(right)
                elif _is_job_title(right):
                    entry["title"]   = right
                    entry["company"] = _strip_city_suffix(left)
                else:
                    # Default: left=title, right=company
                    entry["title"]   = left
                    entry["company"] = _strip_city_suffix(right)

            # Format B/C: no separator — title and company on separate lines
            elif len(header_lines) >= 2:
                line1, line2 = header_lines[0], header_lines[1]
                if _is_job_title(line1) and not _is_job_title(line2):
                    entry["title"]   = line1
                    entry["company"] = _strip_city_suffix(line2)
                elif _is_job_title(line2) and not _is_job_title(line1):
                    entry["title"]   = line2
                    entry["company"] = _strip_city_suffix(line1)
                else:
                    # Both look like titles or neither does — heuristic:
                    # line1 is title, line2 is company
                    entry["title"]   = line1
                    entry["company"] = _strip_city_suffix(line2)

            # Format D: single line, no separator
            else:
                if _is_job_title(first):
                    entry["title"] = first
                else:
                    entry["company"] = _strip_city_suffix(first)

        # ── 4. Description from bullets ────────────────────────────────────
        bullets = []
        for line in desc_lines:
            clean = re.sub(r"^[\u2022\u25aa\u25b8\-\*\u2013\s]+", "", line).strip()
            if clean and not DATE_RANGE_RE.search(clean):
                bullets.append(clean)
        if bullets:
            entry["description"] = " ".join(bullets[:5])

        # ── 5. Only keep if we got at least title, company, or description ──
        if entry["title"] or entry["company"] or entry["description"]:
            entries.append(entry)

    return entries

# ---------------------------------------------------------------------------
# Education extraction
# ---------------------------------------------------------------------------

def _extract_education(edu_text: str) -> List[Dict[str, Any]]:
    """Parse structured education entries."""
    if not edu_text.strip():
        return []

    entries: List[Dict[str, Any]] = []
    
    # Pre-process dense lists
    et = edu_text.strip()
    if "\n\n" not in et:
        lines = et.splitlines()
        new_lines = []
        for i, l in enumerate(lines):
            # Split before lines that have both a date and a comma separator
            if i > 0 and DATE_RANGE_RE.search(l) and "," in l and len(l) < 100:
                new_lines.append("\n" + l)
            else:
                new_lines.append(l)
        et = "\n".join(new_lines)
        
    blocks = re.split(r"\n\s*\n", et)

    for block in blocks:
        block = block.strip()
        if not block:
            continue

        lines = [l.strip() for l in block.splitlines() if l.strip()]
        if not lines:
            continue

        entry: Dict[str, Any] = {
            "institution": None, "degree": None,
            "field_of_study": None,
            "start_date": None, "end_date": None,
        }

        dm = DATE_RANGE_RE.search(block)
        if dm:
            entry["start_date"] = _parse_date(dm.group(1))
            end_raw = dm.group(2)
            if not re.match(r"^(present|current|now)$", end_raw, re.I):
                entry["end_date"] = _parse_date(end_raw)
        else:
            ym = re.search(r"\b(20\d{2}|19\d{2})\b", block)
            if ym:
                entry["end_date"] = f"{ym.group(1)}-01"

        for line in lines:
            deg_m = DEGREE_RE.search(line)
            if deg_m:
                entry["degree"] = deg_m.group(0).strip().rstrip(",")
                after = line[deg_m.end():].strip().lstrip(",-\u2013in. ")
                if after and len(after) < 80:
                    # Remove leading punctuation artifacts like ". Computer Science"
                    after = re.sub(r"^[.\s]+", "", after).strip()
                    # Strip dates from field of study if present
                    dm = DATE_RANGE_RE.search(after)
                    if dm:
                        after = after[:dm.start()].strip()
                    entry["field_of_study"] = after.split(",")[0].strip()
                break

        # Institution: first non-degree, non-date-only, non-CGPA line with meaningful length
        # Scan ALL lines (not just until we find institution) since compact blocks
        # may not have blank line separators
        JUNK_LINE_RE = re.compile(
            r"^\s*(?:cgpa|gpa|grade|percentage|marks|score|result|pass|"
            r"aggregate|division|first|second|merit|distinction)"
            r"|\b\d+\s*/\s*\d+\b"
            r"|\b\d{2,3}\s*%",
            re.I
        )
        # Try lines in order; pick the first clean candidate
        institution_candidates = []
        for line in lines:
            stripped = line.strip()
            # If line has a date, strip it to check if there's an institution name left
            dm = DATE_RANGE_RE.search(stripped)
            if dm:
                stripped = stripped[:dm.start()].strip() + " " + stripped[dm.end():].strip()
                stripped = stripped.strip(" -,\u2013")
            
            if (not DEGREE_RE.search(stripped)
                    and not re.match(r"^\d", stripped)
                    and not JUNK_LINE_RE.search(stripped)
                    and len(stripped) >= 2):   # allow short names like VTU, MIT
                institution_candidates.append(stripped)
        if institution_candidates:
            entry["institution"] = institution_candidates[0]

        if entry["institution"] or entry["degree"]:
            entries.append(entry)

    return entries


# ---------------------------------------------------------------------------
# Projects extraction
# ---------------------------------------------------------------------------

def _extract_projects(proj_text: str) -> List[Dict[str, Any]]:
    """Parse project entries: name, description, technologies, URLs."""
    if not proj_text.strip():
        return []

    projects: List[Dict[str, Any]] = []
    
    # Pre-process to handle dense bulleted lists without blank lines
    pt = proj_text.strip()
    if "\n\n" not in pt:
        # Insert a blank line before any bullet to force separation
        pt = re.sub(r"(\n[ \t]*[\u2022\-\u2013\u25aa\u25b8\*])", r"\n\1", pt)
        
    blocks = re.split(r"\n\s*\n", pt)

    for block in blocks:
        block = block.strip()
        if not block:
            continue
        lines = [l.strip() for l in block.splitlines() if l.strip()]
        if not lines:
            continue

        name = re.sub(r"^[\u2022\-\u2013\u25aa\u25b8\* ]+", "", lines[0]).strip()
        if not name or len(name) > 100:
            continue

        desc_lines = [
            re.sub(r"^[\u2022\-\u2013\u25aa\u25b8\* ]+", "", l).strip()
            for l in lines[1:5]
        ]
        desc = " ".join(d for d in desc_lines if d) or None

        # Find explicit tech line (requires colon — avoids matching description sentences)
        tech: List[str] = []
        for line in lines:
            tm = re.match(
                r"(?:technologies|tech(?:nology)?|built\s+with|stack|tools)"
                r"\s*:\s*(.+)",
                line.strip(), re.I
            )
            if tm:
                tech = [t.strip() for t in re.split(r"[,;]", tm.group(1)) if t.strip()]
                break

        urls = []
        for url in URL_RE.findall(block):
            url = url.rstrip(".,;)")
            if url not in urls:
                urls.append(url)

        projects.append({
            "name":         name,
            "description":  desc,
            "technologies": tech,
            "urls":         urls,
        })

    return projects


# ---------------------------------------------------------------------------
# Certifications extraction
# ---------------------------------------------------------------------------

def _extract_certifications(cert_text: str, full_text: str) -> List[str]:
    """Extract certification names."""
    certs: List[str] = []
    seen: set = set()

    for line in cert_text.splitlines():
        line = re.sub(r"^[\u2022\-\u2013\u25aa\u25b8\* ]+", "", line).strip()
        if line and 5 < len(line) < 120:
            key = line.lower()
            if key not in seen:
                seen.add(key)
                certs.append(line)

    # Fallback: scan full text for cert keywords if no section found
    if not certs:
        for m in CERT_RE.finditer(full_text):
            cert = m.group(0).strip().rstrip(".,;:")
            key = cert.lower()
            if cert and key not in seen:
                seen.add(key)
                certs.append(cert)

    return certs[:10]


# ---------------------------------------------------------------------------
# Confidence helpers
# ---------------------------------------------------------------------------

def _field_confidence(value: Any, base: float, method: str) -> float:
    if not value and value != 0:
        return 0.0
    mult = {
        "structured_json": 1.0,
        "section_parser":  0.90,
        "regex_extraction": 0.80,
        "heuristic":       0.65,
        "nlp_ner":         0.85,
    }
    return round(min(1.0, base * mult.get(method, 0.80)), 4)


# ---------------------------------------------------------------------------
# Signal for JSON fixture redirect
# ---------------------------------------------------------------------------

class _FixtureSignal(Exception):
    def __init__(self, data: dict):
        self.data = data


# ---------------------------------------------------------------------------
# Main adapter class
# ---------------------------------------------------------------------------

class ResumeAdapter(SourceAdapter):
    """
    Hybrid resume adapter.

    Extraction pipeline:
      1. Text extraction  (PDF / DOCX / TXT)
      2. Section detection (dynamic heading scan)
      3. Rule-based parsing per section
      4. NER for name (spaCy if available, heuristic fallback)
      5. Normalization delegated to pipeline runner
      6. Field-level confidence scoring
    """

    TRUST_SCORE   = 0.78
    FIXTURE_TRUST = 0.90

    def can_handle(self, source: Source) -> bool:
        if source.type == "resume":
            return True
        if source.path:
            ext = Path(source.path).suffix.lower()
            return ext in {".pdf", ".docx", ".doc", ".txt"}
        return False

    def get_trust_score(self) -> float:
        return self.TRUST_SCORE

    def extract(self, source: Source) -> RawRecord:
        try:
            if source.content is not None and isinstance(source.content, dict):
                return self._from_fixture(source)
            return self._from_file_or_text(source)
        except _FixtureSignal as sig:
            return self._from_fixture(Source(type="resume", content=sig.data))
        except Exception as exc:
            # Robustness: return partial record rather than crashing
            return RawRecord(
                source="resume",
                source_id=source.path or "",
                fields={
                    "full_name": "", "emails": [], "phones": [],
                    "location": "", "headline": "", "links": [],
                    "skills": [], "experience": [], "education": [],
                    "projects": [], "certifications": [],
                },
                extraction_stats={
                    "format": "error",
                    "error": str(exc),
                    "fields_extracted": 0,
                    "extraction_method": "failed",
                },
            )

    # ------------------------------------------------------------------
    # JSON fixture (pre-structured, high confidence)
    # ------------------------------------------------------------------

    def _from_fixture(self, source: Source) -> RawRecord:
        data: Dict[str, Any] = source.content  # type: ignore[assignment]

        norm_links: List[Dict] = []
        for lk in data.get("links", []):
            if isinstance(lk, dict):
                norm_links.append(lk)
            elif isinstance(lk, str):
                t = "github" if "github.com" in lk else (
                    "linkedin" if "linkedin.com" in lk else "other"
                )
                norm_links.append({"url": lk, "type": t})

        fields = {
            "full_name":      data.get("full_name", ""),
            "emails":         data.get("emails", []),
            "phones":         data.get("phones", []),
            "location":       data.get("location", ""),
            "headline":       data.get("headline", ""),
            "links":          norm_links,
            "skills":         data.get("skills", []),
            "experience":     data.get("experience", []),
            "education":      data.get("education", []),
            "projects":       data.get("projects", []),
            "certifications": data.get("certifications", []),
        }
        return RawRecord(
            source="resume",
            source_id=data.get("id", ""),
            ingested_at=data.get("updated_at", ""),
            fields=fields,
            extraction_stats={
                "format":            "json_fixture",
                "extraction_method": "structured_json",
                "fields_extracted":  sum(1 for v in fields.values() if v),
                "confidence_base":   self.FIXTURE_TRUST,
            },
        )

    # ------------------------------------------------------------------
    # File / text extraction
    # ------------------------------------------------------------------

    def _from_file_or_text(self, source: Source) -> RawRecord:
        text = self._load_text(source)
        if not text.strip():
            raise ValueError("ResumeAdapter: extracted text is empty")
        return self._extract_from_text(text, source)

    def _load_text(self, source: Source) -> str:
        if source.content and isinstance(source.content, str):
            return source.content
        if not source.path:
            raise ValueError("ResumeAdapter: no path or content provided")
        path = Path(source.path)
        ext  = path.suffix.lower()
        if ext == ".pdf":
            return self._extract_pdf(path)
        if ext in {".docx", ".doc"}:
            return self._extract_docx(path)
        if ext == ".txt":
            return path.read_text(encoding="utf-8", errors="replace")
        if ext == ".json":
            data = json.loads(path.read_text(encoding="utf-8"))
            raise _FixtureSignal(data)
        raise ValueError(f"ResumeAdapter: unsupported format '{ext}'")

    def _extract_from_text(self, text: str, source: Source) -> RawRecord:
        with open("debug_resume.txt", "w", encoding="utf-8") as f:
            f.write(text)

        sections     = _split_sections(text)
        header_text  = sections.get("header", text[:600])

        contact      = _extract_contact(header_text, text)
        skills       = _extract_skills(sections.get("skills", ""), "")
        # Extract only tech keywords from projects, not project names
        proj_tech    = _extract_project_tech(sections.get("projects", ""))
        all_skills   = list({s.lower(): s for s in skills + proj_tech}.values())

        experience     = _extract_experience(sections.get("experience", ""))
        education      = _extract_education(sections.get("education", ""))
        projects       = _extract_projects(sections.get("projects", ""))
        certifications = _extract_certifications(
            sections.get("certifications", ""), text
        )

        # Headline from summary section — skip the section header line itself
        headline = ""
        summary  = sections.get("summary", "")
        if summary:
            for line in summary.strip().splitlines():
                line = line.strip()
                # Skip lines that are just section headers, URLs, or very short
                if not line or len(line) < 10:
                    continue
                if SECTION_HEADER_RE.match(line):
                    continue
                if re.search(r"https?://|@", line):
                    continue
                headline = line[:250].strip()
                break

        base   = self.TRUST_SCORE
        method = "section_parser"
        fields = {
            "full_name":      contact["full_name"],
            "emails":         contact["emails"],
            "phones":         contact["phones"],
            "location":       contact["location"],
            "headline":       headline,
            "links":          contact["links"],
            "skills":         all_skills,
            "experience":     experience,
            "education":      education,
            "projects":       projects,
            "certifications": certifications,
        }

        return RawRecord(
            source="resume",
            source_id=source.path or "",
            fields=fields,
            extraction_stats={
                "format":            (
                    Path(source.path).suffix.lower() if source.path else "text"
                ),
                "text_length":       len(text),
                "sections_detected": [k for k in sections if k != "header"],
                "extraction_method": "hybrid_section_parser",
                "spacy_available":   _spacy_available(),
                "fields_extracted":  sum(1 for v in fields.values() if v),
                "confidence_base":   base,
                "field_confidences": {
                    "full_name":  _field_confidence(
                        contact["full_name"], base,
                        "nlp_ner" if _spacy_available() else "heuristic"
                    ),
                    "emails":     _field_confidence(contact["emails"],  base, "regex_extraction"),
                    "phones":     _field_confidence(contact["phones"],  base, "regex_extraction"),
                    "location":   _field_confidence(contact["location"],base, "regex_extraction"),
                    "skills":     _field_confidence(all_skills,         base, method),
                    "experience": _field_confidence(experience,         base, method),
                    "education":  _field_confidence(education,          base, method),
                },
            },
        )

    # ------------------------------------------------------------------
    # Static text extractors
    # ------------------------------------------------------------------

    @staticmethod
    def _extract_pdf(path: Path) -> str:
        try:
            import pdfplumber
            pages = []
            with pdfplumber.open(path) as pdf:
                for page in pdf.pages:
                    t = page.extract_text(x_tolerance=2, y_tolerance=2)
                    if t:
                        pages.append(t)
            return "\n".join(pages)
        except ImportError:
            raise ImportError("pdfplumber not installed: pip install pdfplumber")

    @staticmethod
    def _extract_docx(path: Path) -> str:
        try:
            from docx import Document
            doc   = Document(str(path))
            parts = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    row_text = "  |  ".join(
                        c.text.strip() for c in row.cells if c.text.strip()
                    )
                    if row_text:
                        parts.append(row_text)
            return "\n".join(parts)
        except ImportError:
            raise ImportError("python-docx not installed: pip install python-docx")
